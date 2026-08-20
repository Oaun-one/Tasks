from pathlib import Path

from pydantic import field_validator

from ..models import StrictModel
from ..source_types import SourceCommand, SourceContext


class DocxExtractTextInput(StrictModel):
    """Input for docx.extract_text.

    Attributes:
        path: Workspace-relative DOCX path.
    """

    path: str

    @field_validator("path")
    @classmethod
    def require_docx_extension(cls, value: str) -> str:
        """Requires the path to target a DOCX file.

        Args:
            value: Workspace-relative path from verifier.json.

        Returns:
            The unchanged path when it has a DOCX extension.

        Raises:
            ValueError: If the path does not end in ``.docx``.
        """
        if Path(value).suffix.lower() != ".docx":
            raise ValueError("docx.extract_text requires a .docx file")
        return value


class DocxExtractTextOutput(StrictModel):
    """Output from docx.extract_text.

    Attributes:
        text: Extracted body paragraph and table text.
    """

    text: str


class ExtractText(SourceCommand[DocxExtractTextInput, DocxExtractTextOutput]):
    """Extracts text from DOCX body paragraphs and tables."""

    name = "extract_text"
    input_model = DocxExtractTextInput
    output_model = DocxExtractTextOutput

    def run(self, source_input: DocxExtractTextInput, context: SourceContext) -> DocxExtractTextOutput:
        """Runs DOCX text extraction.

        Args:
            source_input: Validated DOCX extraction input.
            context: Source runtime context.

        Returns:
            Extracted DOCX text.

        """
        resolved = context.resolve_path(source_input.path)
        return DocxExtractTextOutput(text=extract_docx(resolved, context.max_content_chars))


def extract_docx(path: Path, limit: int) -> str:
    """Extracts body paragraph and table text from a DOCX file.

    Args:
        path: Resolved DOCX file path.
        limit: Maximum number of characters to return.

    Returns:
        Extracted text capped to ``limit`` characters.
    """
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    total = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            total = append_limited(parts, paragraph.text, total, limit)
            if total >= limit:
                return "\n".join(parts)[:limit]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                total = append_limited(parts, " | ".join(cells), total, limit)
                if total >= limit:
                    return "\n".join(parts)[:limit]
    # Headers/footers (letterheads, page numbers) live outside doc.paragraphs.
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in header_footer.paragraphs:
                if paragraph.text.strip():
                    total = append_limited(parts, paragraph.text, total, limit)
                    if total >= limit:
                        return "\n".join(parts)[:limit]
    return "\n".join(parts)[:limit]


def append_limited(parts: list[str], text: str, total: int, limit: int) -> int:
    """Appends text and returns the updated character count.

    Args:
        parts: Text fragments collected so far.
        text: Text fragment to append.
        total: Current approximate character count.
        limit: Maximum desired character count.

    Returns:
        Updated approximate character count.
    """
    parts.append(text)
    return total + len(text) + 1


COMMANDS = (ExtractText(),)
